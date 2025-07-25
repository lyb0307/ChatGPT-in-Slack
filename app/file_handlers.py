import logging
from typing import List, Optional
import PyPDF2
from io import BytesIO, StringIO
import json
import csv

from app.slack_ops import download_slack_file_content
from slack_bolt import BoltContext


def append_file_content_if_exists(
    *,
    bot_token: str,
    files: List[dict],
    content: List[dict],
    logger: logging.Logger,
) -> None:
    """
    Process various file types and append their content to the message.
    
    Args:
        bot_token: Slack bot token for downloading files
        files: List of Slack file metadata
        content: List to append processed content to
        logger: Logger instance
    """
    if files is None or len(files) == 0:
        return

    logger.debug(f"Processing {len(files)} files for content extraction")
    
    for file in files:
        mime_type = file.get("mimetype", "")
        filename = file.get("name", "")
        file_url = file.get("url_private")
        
        if not file_url:
            logger.warning(f"No private URL for file: {filename}")
            continue
        
        logger.debug(f"Checking file: {filename}, mime_type: {mime_type}")
        
        try:
            # Download file
            file_bytes = download_slack_file_content(file_url, bot_token)
            
            # Process based on file type
            text_content = None
            
            # PDF files
            if mime_type == "application/pdf" or filename.lower().endswith(".pdf"):
                text_content = extract_pdf_content(file_bytes, filename, logger)
            
            # Word documents
            elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                             "application/msword"] or filename.lower().endswith((".docx", ".doc")):
                text_content = extract_word_content(file_bytes, filename, logger)
            
            # Excel files
            elif mime_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                             "application/vnd.ms-excel"] or filename.lower().endswith((".xlsx", ".xls")):
                text_content = extract_excel_content(file_bytes, filename, logger)
            
            # CSV files
            elif mime_type == "text/csv" or filename.lower().endswith(".csv"):
                text_content = extract_csv_content(file_bytes, filename, logger)
            
            # JSON files
            elif mime_type == "application/json" or filename.lower().endswith(".json"):
                text_content = extract_json_content(file_bytes, filename, logger)
            
            # Plain text and code files
            elif mime_type.startswith("text/") or is_code_file(filename):
                text_content = extract_text_content(file_bytes, filename, logger)
            
            # Add extracted content to message
            if text_content:
                content.append({
                    "type": "text",
                    "text": text_content
                })
                logger.info(f"Successfully processed file: {filename}")
            else:
                logger.info(f"No text content extracted from: {filename}")
                
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            continue


def extract_pdf_content(file_bytes: bytes, filename: str, logger: logging.Logger) -> Optional[str]:
    """Extract text content from PDF files."""
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        text_parts = []
        
        num_pages = len(pdf_reader.pages)
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text().strip()
                if page_text:
                    text_parts.append(f"[Page {page_num + 1} of {num_pages}]")
                    text_parts.append(page_text)
                    text_parts.append("")  # Empty line between pages
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1} of {filename}: {e}")
                continue
        
        if not text_parts:
            return None
        
        full_text = "\n".join(text_parts)
        return format_file_content(full_text, filename, "PDF Document")
        
    except Exception as e:
        logger.error(f"Failed to read PDF {filename}: {e}")
        return None


def extract_word_content(file_bytes: bytes, filename: str, logger: logging.Logger) -> Optional[str]:
    """Extract text content from Word documents (.doc and .docx)."""
    try:
        # First try with python-docx for .docx files (better formatting preservation)
        if filename.lower().endswith('.docx'):
            try:
                import docx
                doc = docx.Document(BytesIO(file_bytes))
                text_parts = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract tables
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            table_text.append(" | ".join(row_text))
                    if table_text:
                        text_parts.append("\n[Table]\n" + "\n".join(table_text) + "\n[End Table]\n")
                
                if text_parts:
                    full_text = "\n\n".join(text_parts)
                    return format_file_content(full_text, filename, "Word Document")
            except Exception as e:
                logger.debug(f"Failed to read .docx with python-docx, trying docx2txt: {e}")
        
        # For .doc files or when python-docx fails, try alternative methods
        if filename.lower().endswith('.doc'):
            logger.info(f"Processing legacy .doc file: {filename}")
            
            # Try extracting as RTF or plain text
            try:
                # Check if it might be RTF format (some .doc files are actually RTF)
                if file_bytes.startswith(b'{\\rtf'):
                    logger.debug("File appears to be RTF format")
                    try:
                        import striprtf
                        text = striprtf.rtf_to_text(file_bytes.decode('latin-1', errors='ignore'))
                        if text and text.strip():
                            return format_file_content(text, filename, "Word Document (RTF)")
                    except ImportError:
                        logger.debug("striprtf not available, trying basic extraction")
                
                # Try basic text extraction with encoding detection
                text = None
                for encoding in ['utf-8', 'latin-1', 'cp1252', 'utf-16']:
                    try:
                        text = file_bytes.decode(encoding)
                        # Basic cleanup - remove null bytes and control characters
                        text = text.replace('\x00', '').strip()
                        if text and len(text) > 10:  # Ensure we got meaningful content
                            logger.debug(f"Successfully decoded .doc file with {encoding} encoding")
                            # Try to clean up binary artifacts
                            import re
                            text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)
                            return format_file_content(text, filename, f"Word Document (Legacy, {encoding})")
                    except UnicodeDecodeError:
                        continue
                
                if not text:
                    logger.warning(f"Could not extract text from legacy .doc file: {filename}. File may be in binary format.")
                    return format_file_content(
                        "Unable to extract text from this legacy .doc file. The file may be in an old binary format that requires Microsoft Word to open.", 
                        filename, 
                        "Word Document (Unsupported Format)"
                    )
                    
            except Exception as e:
                logger.error(f"Error processing .doc file {filename}: {e}")
                return None
        
        # Try with docx2txt for .docx files as fallback
        if filename.lower().endswith('.docx'):
            try:
                import docx2txt
                # Save to temporary file as docx2txt requires file path
                import tempfile
                import os
                
                # Create temp file with proper extension
                suffix = os.path.splitext(filename)[1]
                with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp_file:
                    tmp_file.write(file_bytes)
                    tmp_file_path = tmp_file.name
                
                try:
                    text = docx2txt.process(tmp_file_path)
                    # Clean up temporary file
                    os.unlink(tmp_file_path)
                    
                    if text and text.strip():
                        return format_file_content(text, filename, "Word Document")
                    else:
                        logger.warning(f"No text content extracted from Word document: {filename}")
                        return None
                except Exception as e:
                    # Clean up temporary file on error
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)
                    logger.error(f"docx2txt failed for {filename}: {e}")
                    return None
                    
            except ImportError:
                logger.error("docx2txt not installed. Install with: pip install docx2txt")
                return None
            
    except Exception as e:
        logger.error(f"Failed to read Word document {filename}: {e}")
        return None


def extract_excel_content(file_bytes: bytes, filename: str, logger: logging.Logger) -> Optional[str]:
    """Extract content from Excel files."""
    try:
        import openpyxl
        
        workbook = openpyxl.load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
        text_parts = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = []
            
            # Get max row and column with data
            max_row = min(sheet.max_row, 1000)  # Limit to first 1000 rows
            max_col = min(sheet.max_column, 50)  # Limit to first 50 columns
            
            for row in sheet.iter_rows(min_row=1, max_row=max_row, max_col=max_col, values_only=True):
                # Filter out empty rows
                if any(cell is not None for cell in row):
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    sheet_data.append(" | ".join(row_text))
            
            if sheet_data:
                text_parts.append(f"[Sheet: {sheet_name}]")
                text_parts.extend(sheet_data[:500])  # Limit rows per sheet
                if len(sheet_data) > 500:
                    text_parts.append(f"... ({len(sheet_data) - 500} more rows)")
                text_parts.append("")
        
        if not text_parts:
            return None
        
        full_text = "\n".join(text_parts)
        return format_file_content(full_text, filename, "Excel Spreadsheet")
        
    except ImportError:
        logger.error("openpyxl not installed. Install with: pip install openpyxl")
        return None
    except Exception as e:
        logger.error(f"Failed to read Excel file {filename}: {e}")
        return None


def extract_csv_content(file_bytes: bytes, filename: str, logger: logging.Logger) -> Optional[str]:
    """Extract content from CSV files."""
    try:
        # Try different encodings
        text = None
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                text = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            logger.error(f"Failed to decode CSV file {filename}")
            return None
        
        csv_reader = csv.reader(StringIO(text))
        rows = []
        
        for i, row in enumerate(csv_reader):
            if i >= 1000:  # Limit to first 1000 rows
                rows.append(f"... ({sum(1 for _ in csv_reader)} more rows)")
                break
            if row:  # Skip empty rows
                rows.append(" | ".join(row))
        
        if not rows:
            return None
        
        full_text = "\n".join(rows)
        return format_file_content(full_text, filename, "CSV File")
        
    except Exception as e:
        logger.error(f"Failed to read CSV file {filename}: {e}")
        return None


def extract_json_content(file_bytes: bytes, filename: str, logger: logging.Logger) -> Optional[str]:
    """Extract and format JSON content."""
    try:
        text = file_bytes.decode('utf-8')
        # Parse and pretty-print JSON
        json_data = json.loads(text)
        formatted_json = json.dumps(json_data, indent=2, ensure_ascii=False)
        
        return format_file_content(formatted_json, filename, "JSON File")
        
    except json.JSONDecodeError as e:
        logger.error(f"Invalid JSON in file {filename}: {e}")
        # Return raw text if JSON is invalid
        try:
            text = file_bytes.decode('utf-8')
            return format_file_content(text, filename, "JSON File (Invalid)")
        except:
            return None
    except Exception as e:
        logger.error(f"Failed to read JSON file {filename}: {e}")
        return None


def extract_text_content(file_bytes: bytes, filename: str, logger: logging.Logger) -> Optional[str]:
    """Extract content from plain text and code files."""
    try:
        # Try different encodings
        text = None
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                text = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            logger.error(f"Failed to decode text file {filename}")
            return None
        
        # Determine file type for better formatting
        file_type = "Text File"
        if is_code_file(filename):
            ext = filename.lower().split('.')[-1] if '.' in filename else ''
            file_type = f"Code File ({ext})"
        elif filename.lower().endswith('.md'):
            file_type = "Markdown File"
        
        return format_file_content(text, filename, file_type)
        
    except Exception as e:
        logger.error(f"Failed to read text file {filename}: {e}")
        return None


def is_code_file(filename: str) -> bool:
    """Check if a file is a code file based on extension."""
    code_extensions = {
        'py', 'js', 'ts', 'jsx', 'tsx', 'java', 'cpp', 'c', 'h', 'hpp', 'cs',
        'rb', 'go', 'rs', 'swift', 'kt', 'php', 'sql', 'sh', 'bash', 'yml', 
        'yaml', 'xml', 'html', 'css', 'scss', 'sass', 'r', 'scala', 'lua',
        'pl', 'pm', 'dart', 'vim', 'dockerfile', 'makefile'
    }
    
    if '.' not in filename:
        return filename.lower() in {'dockerfile', 'makefile'}
    
    ext = filename.lower().split('.')[-1]
    return ext in code_extensions


def format_file_content(text: str, filename: str, file_type: str, max_chars: int = 50000) -> str:
    """Format file content with proper headers and truncation."""
    if len(text) > max_chars:
        text = text[:max_chars] + f"\n\n[Content truncated - showing first {max_chars} characters of {len(text)} total]"
    
    return f"[{file_type}: {filename}]\n\n{text}\n\n[End of {file_type}: {filename}]"


def can_process_files(context: BoltContext) -> bool:
    """Check if file processing is enabled."""
    from app.slack_ops import can_access_slack_files
    return can_access_slack_files(context)